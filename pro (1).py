import tkinter as tk
from tkinter import messagebox
import mysql.connector
from docx import Document

# MySQL Database connection
def con1():
    return mysql.connector.connect(
        host="localhost",  # Change to your host if different
        user="root",       # MySQL username
        password="root",   # MySQL password
        database="storemag"
    )

# Function to fetch products from the database
def get_products_by_category(category=""):
    connection = con1()
    cursor = connection.cursor()
    query = "SELECT * FROM product WHERE category LIKE %s" if category else "SELECT * FROM product"
    cursor.execute(query, ('%' + category + '%',) if category else ())
    products = cursor.fetchall()
    connection.close()
    return products

# Store Panel to show products category-wise
def storepanel():
    window.configure(bg="light green")
    clear_screen()
    label_info = tk.Label(window, text="Store Management", font=("Arial", 12), fg="blue")
    label_info.pack(pady=20)

    # Category selection
    label_category = tk.Label(window, text="Enter Category (optional):", fg="red")
    label_category.pack(pady=5)

    entry_category = tk.Entry(window)
    entry_category.pack(pady=5)

    # Search button to filter products by category
    button_search = tk.Button(window, text="Search", command=lambda: display_products(entry_category.get()), bg="pink")
    button_search.pack(pady=5)

    # Back button
    button_back = tk.Button(window, text="Back", command=show_main_panel, bg="red", fg="white")
    button_back.pack(pady=10)

def display_products(category):
    clear_screen()
    label_category = tk.Label(window, text="Store Management", font=("Arial", 12), fg="blue")
    label_category.pack(pady=20)

    products = get_products_by_category(category)

    cart_items = {}

    def update_cart_display():
        cart_display.delete(0, tk.END)
        for name, qty in cart_items.items():
            cart_display.insert(tk.END, f"{name} - Qty: {qty}")

    def add_to_cart(product_id, name, price, available_qty):
        if name in cart_items:
            if cart_items[name] < available_qty:
                cart_items[name] += 1
            else:
                messagebox.showwarning("Stock Limit", f"Cannot add more than available quantity ({available_qty}) for {name}.")
        else:
            if available_qty > 0:
                cart_items[name] = 1
            else:
                messagebox.showwarning("Out of Stock", f"{name} is out of stock.")
        update_cart_display()

    def remove_from_cart(product_id, name):
        if name in cart_items and cart_items[name] > 0:
            cart_items[name] -= 1
        if cart_items[name] == 0:
            del cart_items[name]
        update_cart_display()

    # Display products with + and - buttons
    for product in products:
        product_id, category, name, qty, mrp = product

        # Create a horizontal frame for each product
        product_frame = tk.Frame(window)
        product_frame.pack(fill=tk.X, padx=10, pady=5)

        # Product information
        product_label = tk.Label(product_frame, text=f"{name} - ${mrp} (Available: {qty})", width=50, anchor="w", fg="blue")
        product_label.pack(side=tk.LEFT, padx=5)

        # Add (+) button
        button_add = tk.Button(product_frame, text="+", command=lambda pid=product_id, pname=name, pprice=mrp, pqty=qty: add_to_cart(pid, pname, pprice, pqty), bg="green", fg="white")
        button_add.pack(side=tk.LEFT, padx=5)

        # Remove (-) button
        button_remove = tk.Button(product_frame, text="-", command=lambda pid=product_id, pname=name: remove_from_cart(pid, pname), bg="orange", fg="white")
        button_remove.pack(side=tk.LEFT, padx=5)

    # Cart display
    cart_display = tk.Listbox(window, width=50, height=10)
    cart_display.pack(pady=20)

    # Proceed to bill button
    button_proceed_to_bill = tk.Button(window, text="Proceed to Bill", command=lambda: billpanel(cart_items), bg="purple", fg="white")
    button_proceed_to_bill.pack(pady=10)

    # Back button
    button_back = tk.Button(window, text="Back", command=show_main_panel, bg="red", fg="white")
    button_back.pack(pady=10)

# Bill Panel for customer details and bill generation
def billpanel(cart_items):
    clear_screen()
    window.configure(bg="light blue")

    label_info = tk.Label(window, text="Bill Generation", font=("Arial", 14), fg="blue")
    label_info.pack(pady=20)

    total_price = 0
    bill_details = ""

    # Fetching the actual MRP and calculating the total
    connection = con1()
    cursor = connection.cursor()

    for item_name, qty in cart_items.items():
        cursor.execute("SELECT mrp FROM product WHERE prodnm = %s", (item_name,))
        result = cursor.fetchone()
        if result:
            mrp = result[0]
            total_price += qty * mrp
            bill_details += f"{item_name} - Qty: {qty}  ${mrp} each\n"
        else:
            bill_details += f"{item_name} - Qty: {qty} (Price not found)\n"

    connection.close()

    label_items = tk.Label(window, text="Selected Items:\n" + bill_details, fg="blue")
    label_items.pack()

    label_total = tk.Label(window, text=f"Total: ${total_price}", fg="green")
    label_total.pack()

    # Customer details entry fields
    label_name = tk.Label(window, text="Name:", fg="red")
    label_name.pack()
    entry_name = tk.Entry(window)
    entry_name.pack(pady=5)

    label_phone = tk.Label(window, text="Phone Number:", fg="red")
    label_phone.pack()
    entry_phone = tk.Entry(window)
    entry_phone.pack(pady=5)

    label_payment = tk.Label(window, text="Mode of Payment:", fg="red")
    label_payment.pack()
    entry_payment = tk.Entry(window)
    entry_payment.pack(pady=5)

    def save_bill():
        name = entry_name.get()
        phone = entry_phone.get()
        payment_mode = entry_payment.get()

        if not name or not phone or not payment_mode:
            messagebox.showerror("Input Error", "Please fill all the customer details.")
            return

        try:
            conn = con1()
            cursor = conn.cursor()

            # Update the stock in the database
            for item_name, qty in cart_items.items():
                cursor.execute("UPDATE product SET qty = qty - %s WHERE prodnm = %s AND qty >= %s", (qty, item_name, qty))
                if cursor.rowcount == 0:
                    messagebox.showwarning("Stock Warning", f"Insufficient stock for {item_name}. Stock not updated.")

            # Add the sale to the sales table
            product_details = "\n".join([f"{item}: {qty}" for item, qty in cart_items.items()])
            cursor.execute(
                "INSERT INTO sales (customer_name, phone_number, payment_mode, product_details, total_price) VALUES (%s, %s, %s, %s, %s)",
                (name, phone, payment_mode, product_details, total_price)
            )

            conn.commit()
            conn.close()

            # Create the Word file
            doc = Document()
            doc.add_heading('Store Bill', 0)
            doc.add_paragraph(f"Customer Name: {name}")
            doc.add_paragraph(f"Phone No: {phone}")
            doc.add_paragraph(f"Mode of Payment: {payment_mode}")
            doc.add_paragraph(f"\nSelected Products:\n{bill_details}")
            doc.add_paragraph(f"\nTotal Price: ${total_price}")

            # Save the document
            doc.save(f"{name}_bill.docx")
            messagebox.showinfo("Success", "Bill saved successfully. Stock updated.")

            # Redirect to main panel after saving bill
            show_main_panel()

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

    button_save = tk.Button(window, text="Save Bill", command=save_bill, bg="green", fg="white")
    button_save.pack(pady=5)

    # Back button
    button_back = tk.Button(window, text="Back", command=storepanel, bg="red", fg="white")
    button_back.pack(pady=10)

# Show Sales Panel to display sales records date-wise



# Function to add a product to the database
def add_product():
    category = entry_category.get()
    prodnm = entry_prodnm.get()
    qty = entry_qty.get()
    mrp = entry_mrp.get()

    if not category or not prodnm or not qty or not mrp:
        messagebox.showerror("Input Error", "Please fill all fields.")
        return

    try:
        conn = con1()
        cursor = conn.cursor()
        cursor.execute("INSERT INTO product (category, prodnm, qty, mrp) VALUES (%s, %s, %s, %s)", 
                       (category, prodnm, qty, mrp))
        conn.commit()
        conn.close()
        messagebox.showinfo("Success", "Product added successfully.")
        clear_entries()
    except Exception as e:
        messagebox.showerror("Database Error", str(e))

# Function to remove a product from the database
def remove_product():
    prodnm = entry_prodnm.get()
    if not prodnm:
        messagebox.showerror("Input Error", "Please provide a product name.")
        return

    try:
        conn = con1()
        cursor = conn.cursor()
        cursor.execute("DELETE FROM product WHERE prodnm = %s", (prodnm,))
        conn.commit()
        conn.close()
        messagebox.showinfo("Success", "Product removed successfully.")
        clear_entries()
    except Exception as e:
        messagebox.showerror("Database Error", str(e))

# Function to modify a product in the database
def modify_product():
    prodnm = entry_prodnm.get()
    qty = entry_qty.get()
    mrp = entry_mrp.get()

    if not prodnm or not qty or not mrp:
        messagebox.showerror("Input Error", "Please provide all the required details.")
        return

    try:
        conn = con1()
        cursor = conn.cursor()
        cursor.execute("UPDATE product SET qty = %s, mrp = %s WHERE prodnm = %s", 
                       (qty, mrp, prodnm))
        conn.commit()
        conn.close()
        messagebox.showinfo("Success", "Product updated successfully.")
        clear_entries()
    except Exception as e:
        messagebox.showerror("Database Error", str(e))

# Function to display the Admin panel
def adminpanel():
    window.configure(bg="light pink")
    clear_screen()
    global label_category,label_prodnm,label_qty,label_mrp,entry_category,entry_prodnm,entry_prodnm,entry_prodnm,button_remove,button_add,button_modify	

    label_category.pack()
    entry_category.pack()
    label_prodnm.pack()
    entry_prodnm.pack()
    label_qty.pack()
    entry_qty.pack()
    label_mrp.pack()
    entry_mrp.pack()

    button_add.pack()
    button_remove.pack()
    button_modify.pack()

    # Back button
    button_back = tk.Button(window, text="Back", command=show_main_panel, bg="red", fg="white")
    button_back.pack(pady=10)

# Function to clear the entry fields
def clear_entries():
    entry_category.delete(0, tk.END)
    entry_prodnm.delete(0, tk.END)
    entry_qty.delete(0, tk.END)
    entry_mrp.delete(0, tk.END)

# Sales Panel for displaying sales records date-wise
def salespanel():
    clear_screen()
    window.configure(bg="#add8e6")

    label_info = tk.Label(window, text="Sales Records", font=("Arial", 14), fg="blue")
    label_info.pack(pady=20)

    # Label and Entry for date input
    label_date = tk.Label(window, text="Enter Date (YYYY-MM-DD):", fg="red")
    label_date.pack(pady=5)

    entry_date = tk.Entry(window)
    entry_date.pack(pady=5)

    sales_display = tk.Listbox(window, width=100, height=20)
    sales_display.pack(pady=20)

    def fetch_sales():
        sales_display.delete(0, tk.END)
        date = entry_date.get()

        if not date:
            messagebox.showerror("Input Error", "Please enter a valid date.")
            return

        try:
            conn = con1()
            cursor = conn.cursor()

            # Fetch sales records for the given date
            query = "SELECT * FROM sales WHERE DATE(sale_date) =%s LIMIT 100000"
            cursor.execute(query, (date,))
            sales = cursor.fetchall()
            print(sales)
            conn.close()

            if not sales:
                sales_display.insert(tk.END, "No sales records found for the selected date.")
            else:
                # Display sales records
                for sale in sales:
                    print(sale)
                    sale_id, customer_name, phone_number, payment_mode, product_details, total_price, sale_date = sale
                    sales_display.insert(tk.END, f"Sale ID: {sale_id} | Customer: {customer_name} | Phone: {phone_number} | Payment: {payment_mode} | Total: ${total_price} | Date: {sale_date}")
                    sales_display.insert(tk.END, f"  Products: {product_details}")
                    sales_display.insert(tk.END, "")

        except Exception as e:
            messagebox.showerror("Database Error", f"An error occurred: {e}")

    # Fetch sales button
    button_fetch_sales = tk.Button(window, text="Fetch Sales", command=fetch_sales, bg="green", fg="white")
    button_fetch_sales.pack(pady=10)

    # Back button
    button_back = tk.Button(window, text="Back", command=show_main_panel, bg="red", fg="white")
    button_back.pack(pady=10)

# Function to clear the screen
def clear_screen():
    for widget in window.winfo_children():
        widget.pack_forget()

# Main Panel
def show_main_panel():
    clear_screen()
    title=tk.Label(window,text="PATEL CONTROL PANEL",font=("Arial",30),fg="blue")
    title.pack()
    window.configure(bg="#f0f8ff")
    button_admin.pack(pady=10)
    button_store.pack(pady=10)
    button_bill.pack(pady=10)
    button_sales.pack(pady=10)
    create=tk.Label(window,text="Made By Riddhi and Siddhi",font=("Arial",10),fg="blue")
    create.pack()

# Main GUI setup
window = tk.Tk()
window.title("Store Management System")
  # Set background color

# Labels for inputs
global label_category,label_prodnm,label_qty,label_mrp,entry_category,entry_prodnm,entry_prodnm,entry_prodnm,button_remove,button_add,button_modify,button_admin,button_store,button_bill

label_category = tk.Label(window, text="Category", bg="#f0f8ff", font=("Arial", 10), fg="blue")
label_prodnm = tk.Label(window, text="Product Name", bg="#f0f8ff", font=("Arial", 10), fg="blue")
label_qty = tk.Label(window, text="Quantity", bg="#f0f8ff", font=("Arial", 10), fg="blue")
label_mrp = tk.Label(window, text="MRP", bg="#f0f8ff", font=("Arial", 10), fg="blue")

# Entry fields for inputs
entry_category = tk.Entry(window)
entry_prodnm = tk.Entry(window)
entry_qty = tk.Entry(window)
entry_mrp= tk.Entry(window)

# Buttons for Admin operations
button_add = tk.Button(window, text="Add Product", command=add_product, bg="green", fg="white")
button_remove = tk.Button(window, text="Remove Product", command=remove_product, bg="orange", fg="white")
button_modify = tk.Button(window, text="Modify Product", command=modify_product, bg="blue", fg="white")

# Buttons to switch between panels
button_admin = tk.Button(window, text="Admin", command=adminpanel, bg="red", fg="white",height=3, width=10)  # Admin panel placeholder
button_store = tk.Button(window, text="Store", command=storepanel, bg="purple", fg="white",height=3, width=10)
button_bill = tk.Button(window, text="Bill", command= billpanel, bg="pink", fg="white",height=3, width=10)
button_sales=tk.Button(window,text="sales",command=salespanel,bg='orange',fg='white',height=3, width=10)

# Layout for the initial screen
button_admin.pack(pady=10)
button_store.pack(pady=10)
button_bill.pack(pady=10)
button_sales.pack(pady=10)


# Start at the main panel
show_main_panel()

# Start the Tkinter event loop
window.mainloop()


